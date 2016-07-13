import ConfigParser
import datetime
import os
import ApiClient
import printProgress
import VaultService
import json
import ast
import logging
from docx import Document
import collections
from BeautifulSoup import BeautifulSoup as bs

def process_json(json_object, type, table, json_str, valid):
    # Include full JSON
    row_cells = table.add_row().cells
    
    if valid:
        row_cells[0].text = 'JSON definition'
        row_cells[1].text = json.dumps(json_object, indent=2)
    else:                
        row_cells[0].text = 'JSON definition (invalid)'
        row_cells[1].text = json_str or ''

    if json_object is not None:
        obj = json_object
        if type == 'Pagelayout': #hack until pagelayout fixed
            obj = json_object[0]
        
        write_items(obj, table)     
                        
def write_items(dictionary, table):    
    
        for attr, value in dictionary.iteritems():
            if isinstance(value, collections.Iterable) and not isinstance(value, basestring):
                
                length = len(value)
                row_cells = table.add_row().cells
                row_cells[0].text = write_value(attr).upper()
                
                if len(value) > 0: 
                    i = 1
                    for item in value:
                        row_cells = table.add_row().cells
                        row_cells[0].text = "#" + str(i)
                        
                        # if list of strings write at header row
                        if isinstance(item, basestring):
                            row_cells[1].text = write_value(item)
                        
                        # if json iterate through sub attributes
                        if isinstance(item, collections.Mapping):
                            
                            for attr2, value2 in item.iteritems():
                                row_cells = table.add_row().cells
                                row_cells[0].text = "\t" + str(attr2)
                                row_cells[1].text = write_value(value2)
                        i+=1                        
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = str(attr)
                row_cells[1].text = write_value(value)

def is_json(myjson):
    try:
        json_object = json.loads(myjson)
    except (ValueError, TypeError):
        return False

    if type(json_object) is int:
        return False
    return json_object


def print_json(obj, tbl, parent=None):
    if isinstance(obj, list):
        my_dict = {}

        for idx, val in enumerate(obj):
            my_dict[idx] = val

        obj = my_dict

    for attr, value in obj.iteritems():

        json_obj = is_json(str(value))
        if json_obj is False:
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(attr)
            row_cells[1].text = str(value)
        else:
            print json_obj
            print_json(json_obj, tbl, str(attr))


def can_decode(type, attribute):
    if type == "Doctype" and attribute == 'document_name_format':
        return False
    if type == "Docfield" and attribute == 'formula':
        return False

    return True

def write_value(value):
    if str(value).startswith("<"):
        soup=bs(str(value))                #make BeautifulSoup
        return soup.prettify()   #prettify the html
        # try:
        #     this_xml = xml.dom.minidom.parseString(str(value))
        #     return this_xml.toprettyxml(indent="  ")
        # except xml.parsers.expat.ExpatError:
        #     print "failed XML parsing for " + str(value) 
        #     pass
    
    return str(value)


def process_components(data, type):

    document = Document('../output/config.docx')
    document.add_heading('Config report for %s' % client.domain, 0)

    last_type = None

    for component in data:
        name = component["component_name__v"]
        type = component["component_type__v"]
        checksum = component["checksum__v"]
        mdl = component["mdl_definition__v"]
        json_str = component["json_definition__v"]

        valid = True
        
        # parse string to object
        try:
            json_object = json.loads(json_str)
        except (ValueError,TypeError):
            json_object = None
            valid = False

        if create_doc:
            if (last_type != type):
                last_type = type
                # document.add_page_break()
                document.add_heading(type, level=1)
            
            document.add_heading(name, level=2)

            table = document.add_table(rows=3, cols=2)
            table.style = 'Veeva_Table'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Attribute'
            hdr_cells[1].text = 'Value'

            hdr_cells = table.rows[1].cells
            hdr_cells[0].text = 'valid?'
            hdr_cells[1].text = str(valid)
            
            hdr_cells = table.rows[2].cells
            hdr_cells[0].text = 'checksum'
            hdr_cells[1].text = checksum
            
            if use_format.upper() == "JSON":
                process_json(json_object, type, table, json_str, valid)
            
            if use_format.upper() == "MDL":
                row_cells = table.add_row().cells
                row_cells[0].text = 'MDL'
                row_cells[1].text = mdl        
        
        if create_files:
            folder_name = 'valid'
            if valid is False:
                folder_name = 'invalid'

            type_folder = "../output/JSON/%s/%s/%s/%s" % (client.domain, instance_name, folder_name, type)
            
            if not os.path.exists(type_folder):
                os.makedirs(type_folder)
            
            with open(type_folder + "/" + name + ".json", "w") as f:
                if valid is False:
                    if json_str:
                        f.write(json_str.encode('utf-8'))
                else:
                    json.dump(json_object, f, indent=4)            

    if create_doc:
        document.save('../output/DOCS/%s-%s.docx' % (client.domain, instance_name))

print """
  .oooooo.                          .o88o.  o8o                  oooooooooo.                                          .oooo.   
 d8P'  `Y8b                         888 `"  `"'                  `888'   `Y8b                                       .dP""Y88b  
888           .ooooo.  ooo. .oo.   o888oo  oooo   .oooooooo       888      888  .ooooo.   .ooooo.       oooo    ooo       ]8P' 
888          d88' `88b `888P"Y88b   888    `888  888' `88b        888      888 d88' `88b d88' `"Y8       `88.  .8'      .d8P'  
888          888   888  888   888   888     888  888   888        888      888 888   888 888              `88..8'     .dP'     
`88b    ooo  888   888  888   888   888     888  `88bod8P'        888     d88' 888   888 888   .o8         `888'    .oP     .o 
 `Y8bood8P'  `Y8bod8P' o888o o888o o888o   o888o `8oooooo.       o888bood8P'   `Y8bod8P' `Y8bod8P'          `8'     8888888888 
                                                 d"     YD                                                                     
                                                 "Y88888P'                                                                     
"""

client = VaultService.get_client()

#input = raw_input("D for document, F for files, B for both")
#create_doc = input == 'D'
#create_files = input == 'F'
#if(input == 'B'):
create_doc = create_files = True
use_format = "JSON"

instance_name = datetime.datetime.now()
print "Format used to get components: " + use_format

limit = 1000
size = 1000
offset = 0
i = 1

while size == limit:
    directory = client.get_json("query/components?q=SELECT component_name__v, checksum__v, component_type__v, json_definition__v, mdl_definition__v FROM vault_component__v ORDER BY component_type__v LIMIT %s OFFSET %s" % (limit, offset));
    limit = directory["responseDetails"]["limit"]
    size = directory["responseDetails"]["size"]
    offset += limit
    process_components(directory["data"], use_format)
    print "Batch %s of %s" %(i, size)
    i+=1

client = None
print "Done"
logging.info("-------------------DONE------------------")





